"""Escreve o Markdown-com-matemática do Gerador dentro de um documento do Word.

A fórmula vira **equação nativa** (OMML), não uma imagem nem `$x^2$` em texto:
o professor abre o `.docx`, clica na fórmula e edita. O caminho é
LaTeX → MathML (`latex2mathml`) → OMML (`mathml2omml`), tudo em Python puro —
sem Pandoc, sem Node, sem LaTeX instalado. `instalar.bat` continua sendo só
`pip install`.

As duas bibliotecas são opcionais de propósito. Se faltarem, ou se o LLM emitir
uma fórmula que elas não sabem ler, a lista **continua saindo**: a fórmula cai
para texto, sem os cifrões. Uma lista com a fórmula feia é aproveitável; uma
exceção no meio da exportação não é.
"""

from __future__ import annotations

from .marcacao import Bloco, Trecho, analisar

_ESPACO_OOXML = (
    'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
)


def _equacao(latex: str):
    """O LaTeX como nós OMML, ou `None` quando não dá para converter."""
    try:
        import latex2mathml.converter
        import mathml2omml
        from lxml import etree
    except ImportError:
        return None
    try:
        omml = mathml2omml.convert(latex2mathml.converter.convert(latex))
        return list(etree.fromstring(f"<raiz {_ESPACO_OOXML}>{omml}</raiz>"))
    except Exception:
        # O Gerador escreve LaTeX livre, e nem tudo que o pdflatex aceita as
        # bibliotecas aceitam. Cair para texto é a resposta certa aqui.
        return None


def _escrever_trechos(paragrafo, trechos: list[Trecho]) -> None:
    for t in trechos:
        if t.tipo == "negrito":
            paragrafo.add_run(t.conteudo).bold = True
        elif t.tipo == "matematica":
            nos = _equacao(t.conteudo)
            if nos is None:
                paragrafo.add_run(t.conteudo)
            else:
                for no in nos:
                    paragrafo._p.append(no)
        else:
            paragrafo.add_run(t.conteudo)


def _tabela(doc, bloco: Bloco) -> None:
    tabela = doc.add_table(rows=len(bloco.celulas), cols=len(bloco.celulas[0]))
    tabela.style = "Table Grid"
    for i, linha in enumerate(bloco.celulas):
        for j, celula in enumerate(linha):
            alvo = tabela.cell(i, j).paragraphs[0]
            _escrever_trechos(alvo, celula)
            if i == 0:  # a primeira linha da tabela de faixas é o cabeçalho
                for run in alvo.runs:
                    run.bold = True


def escrever(doc, texto: str, prefixo: str = "", recuo=None) -> None:
    """Põe `texto` no documento, bloco a bloco, com o `prefixo` em negrito."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    blocos = analisar(texto)
    pendente = prefixo

    def novo_paragrafo(estilo: str | None = None):
        nonlocal pendente
        p = doc.add_paragraph(style=estilo)
        if recuo is not None:
            p.paragraph_format.left_indent = recuo
            p.paragraph_format.space_after = Pt(2)
        if pendente:
            p.add_run(pendente).bold = True
            pendente = ""
        return p

    for bloco in blocos:
        if bloco.tipo == "tabela":
            if pendente:
                novo_paragrafo()
            _tabela(doc, bloco)
            doc.add_paragraph()  # respiro entre a tabela e o texto seguinte
        elif bloco.tipo == "lista":
            for item in bloco.linhas:
                _escrever_trechos(novo_paragrafo("List Bullet"), item)
        elif bloco.tipo == "formula":
            p = novo_paragrafo()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _escrever_trechos(p, [Trecho("matematica", bloco.formula, True)])
        else:
            p = novo_paragrafo()
            for n, linha in enumerate(bloco.linhas):
                if n:
                    # Quebra de linha de verdade: o `\n` dentro de um run some,
                    # e era isso que colava a), b) e c) numa linha só.
                    p.add_run().add_break()
                _escrever_trechos(p, linha)

    if pendente:  # texto vazio: ao menos o rótulo da questão aparece
        novo_paragrafo()

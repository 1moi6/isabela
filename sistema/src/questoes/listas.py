"""Montagem e exportação de listas de exercícios a partir do banco curado.

Exporta em Markdown (uso direto/conversão), LaTeX (impressão via pdflatex) e
Word (.docx, para quem prefere editar antes de imprimir), com versão do aluno
(sem gabarito) e do professor (com resolução).
"""

from __future__ import annotations

from io import BytesIO

from .modelos import Questao
from .tex import PREAMBULO, para_tex
from .word import escrever as _escrever_no_word

_CABECALHO_LATEX = PREAMBULO + "\\begin{document}\n"


def para_markdown(
    titulo: str, questoes: list[Questao], com_gabarito: bool = False
) -> str:
    """Lista em Markdown. `com_gabarito=True` gera a versão do professor.

    Aqui o texto do Gerador vai como veio: ele *é* Markdown, e reescrevê-lo só
    poderia estragá-lo. As linhas em branco é que não são decorativas — sem a
    que separa o enunciado das alternativas, o leitor de Markdown absorve a
    lista dentro do parágrafo anterior e as letras (a)–(d) somem.
    """
    partes = [f"# {titulo}", ""]
    for i, q in enumerate(questoes, start=1):
        partes.append(f"**Questão {i}.** {q.enunciado}")
        if q.alternativas:
            partes.append("")
            for letra, alt in zip("abcd", q.alternativas):
                # Alternativa em mais de uma linha continua no mesmo item: sem o
                # recuo, a segunda linha fecha a lista.
                texto = alt.texto.replace("\n", "\n  ")
                partes.append(f"- ({letra}) {texto}")
        partes.append("")
    if com_gabarito:
        partes += ["---", "", "## Gabarito e resoluções", ""]
        for i, q in enumerate(questoes, start=1):
            partes.append(f"**Questão {i}.** {q.gabarito}")
            partes.append("")
            partes.append(q.resolucao)
            partes.append("")
    return "\n".join(partes)


def para_latex(titulo: str, questoes: list[Questao], com_gabarito: bool = False) -> str:
    """Lista em LaTeX pronta para pdflatex.

    Todo texto vindo do Gerador passa por `para_tex`. Antes ia cru, e não
    compilava: `R$ 13,00` abre modo matemático, a tabela de tarifa por faixa sai
    como texto cheio de barras verticais e `50%` come o resto da linha. O botão
    "LaTeX" da interface entregava um arquivo que o professor não conseguia usar.
    """
    partes = [_CABECALHO_LATEX, rf"\section*{{{para_tex(titulo)}}}",
              r"\begin{enumerate}[label=\textbf{\arabic*.}]"]
    for q in questoes:
        partes.append(rf"\item {para_tex(q.enunciado)}")
        if q.alternativas:
            partes.append(r"\begin{enumerate}[label=(\alph*)]")
            partes += [rf"\item {para_tex(alt.texto)}" for alt in q.alternativas]
            partes.append(r"\end{enumerate}")
    partes.append(r"\end{enumerate}")
    if com_gabarito:
        partes += [r"\newpage", r"\section*{Gabarito e resoluções}",
                   r"\begin{enumerate}[label=\textbf{\arabic*.}]"]
        for q in questoes:
            partes.append(rf"\item {para_tex(q.gabarito)}")
            partes.append("")
            partes.append(para_tex(q.resolucao))
        partes.append(r"\end{enumerate}")
    partes.append(r"\end{document}")
    return "\n".join(partes)


def para_docx(titulo: str, questoes: list[Questao], com_gabarito: bool = False) -> bytes:
    """Lista em Word (.docx), pronta para editar antes de imprimir.

    Devolve os bytes do arquivo — quem chama decide se grava em disco ou envia
    pelo navegador.

    O texto passa por `marcacao.analisar` pelo mesmo motivo que passa por
    `para_tex` no caminho do LaTeX: antes ia cru, e o professor recebia
    `$$C(x) = \\begin{cases}…$$`, `**Modelando a situação**` e a tabela de
    tarifa como uma fila de barras verticais — tudo numa linha só, porque o Word
    engole a quebra de linha que vem dentro de um run.
    """
    from docx import Document  # dependência opcional: pip install questoes-em[docx]
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(titulo, level=1)

    for i, q in enumerate(questoes, start=1):
        _escrever_no_word(doc, q.enunciado, prefixo=f"Questão {i}. ")
        for letra, alt in zip("abcd", q.alternativas or []):
            _escrever_no_word(doc, alt.texto, prefixo=f"({letra}) ", recuo=Pt(24))

    if com_gabarito:
        doc.add_page_break()
        doc.add_heading("Gabarito e resoluções", level=1)
        for i, q in enumerate(questoes, start=1):
            _escrever_no_word(doc, q.gabarito, prefixo=f"Questão {i}. ")
            _escrever_no_word(doc, q.resolucao)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
